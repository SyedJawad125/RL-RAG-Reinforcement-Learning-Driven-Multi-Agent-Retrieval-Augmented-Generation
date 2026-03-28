# """
# Document Processing Service
# Handles file upload, text extraction, chunking, and vectorization
# """
# import logging
# from typing import Dict, Any, List
# import PyPDF2
# import docx
# import io
# from django.core.files.uploadedfile import UploadedFile
# from django.conf import settings

# logger = logging.getLogger(__name__)


# class DocumentProcessor:
#     """
#     Processes uploaded documents for RAG system.
    
#     Handles:
#     - Text extraction (PDF, DOCX, TXT)
#     - Text chunking
#     - Embedding generation
#     - Vector store insertion
#     """
    
#     def __init__(self, vector_store, embedding_service, chunk_size=None, chunk_overlap=None):
#         """
#         Initialize document processor.
        
#         Args:
#             vector_store: Vector store instance
#             embedding_service: Embedding service instance
#             chunk_size: Size of text chunks
#             chunk_overlap: Overlap between chunks
#         """
#         self.vector_store = vector_store
#         self.embedding_service = embedding_service
#         self.chunk_size = chunk_size or getattr(settings, 'CHUNK_SIZE', 800)
#         self.chunk_overlap = chunk_overlap or getattr(settings, 'CHUNK_OVERLAP', 100)
    
#     async def process_document(self, file: UploadedFile, document_id: str) -> Dict[str, Any]:
#         """
#         Process uploaded document.
        
#         Args:
#             file: Uploaded file
#             document_id: Document UUID
            
#         Returns:
#             Processing result dictionary
#         """
#         # Extract text
#         text = self._extract_text(file)
        
#         if not text or len(text.strip()) < 10:
#             raise ValueError("No extractable text found in document")
        
#         logger.info(f"[Processor] Extracted {len(text)} characters")
        
#         # Chunk text
#         chunks = self._chunk_text(text)
        
#         if not chunks:
#             raise ValueError("No chunks generated from document")
        
#         logger.info(f"[Processor] Created {len(chunks)} chunks")
        
#         # Generate embeddings
#         embeddings = self.embedding_service.embed_texts(chunks)
        
#         # Prepare metadata
#         metadatas = [
#             {
#                 "source": file.name,
#                 "content_type": file.content_type or 'application/octet-stream',
#                 "chunk_index": i,
#                 "document_id": document_id,
#                 "chunk_size": len(chunk)
#             }
#             for i, chunk in enumerate(chunks)
#         ]
        
#         # Generate IDs
#         ids = [f"{document_id}_chunk_{i}" for i in range(len(chunks))]
        
#         # Add to vector store
#         self.vector_store.add_documents(
#             documents=chunks,
#             embeddings=embeddings,
#             metadata=metadatas,
#             ids=ids
#         )
        
#         logger.info(f"[Processor] Added {len(chunks)} chunks to vector store")
        
#         return {
#             "chunks_created": len(chunks),
#             "text_length": len(text),
#             "document_id": document_id
#         }
    
#     def _extract_text(self, file: UploadedFile) -> str:
#         """Extract text from uploaded file"""
#         content = file.read()
        
#         if file.content_type == 'application/pdf' or file.name.endswith('.pdf'):
#             return self._extract_pdf(content)
#         elif file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file.name.endswith('.docx'):
#             return self._extract_docx(content)
#         elif file.content_type == 'text/plain' or file.name.endswith('.txt'):
#             return content.decode('utf-8', errors='ignore')
#         else:
#             raise ValueError(f"Unsupported file type: {file.content_type}")
    
#     def _extract_pdf(self, content: bytes) -> str:
#         """Extract text from PDF"""
#         try:
#             pdf_file = io.BytesIO(content)
#             pdf_reader = PyPDF2.PdfReader(pdf_file)
#             text = ""
#             for page in pdf_reader.pages:
#                 text += page.extract_text() + "\n"
#             return text.strip()
#         except Exception as e:
#             logger.error(f"PDF extraction failed: {e}")
#             raise ValueError(f"Failed to extract PDF: {e}")
    
#     def _extract_docx(self, content: bytes) -> str:
#         """Extract text from DOCX"""
#         try:
#             doc_file = io.BytesIO(content)
#             doc = docx.Document(doc_file)
#             text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
#             return text.strip()
#         except Exception as e:
#             logger.error(f"DOCX extraction failed: {e}")
#             raise ValueError(f"Failed to extract DOCX: {e}")
    
#     def _chunk_text(self, text: str) -> List[str]:
#         """
#         Split text into overlapping chunks.
        
#         Args:
#             text: Input text
            
#         Returns:
#             List of text chunks
#         """
#         if not text:
#             return []
        
#         words = text.split()
#         if len(words) <= self.chunk_size:
#             return [" ".join(words)]
        
#         chunks = []
#         start = 0
        
#         while start < len(words):
#             end = start + self.chunk_size
#             chunk = " ".join(words[start:end])
            
#             if chunk.strip():
#                 chunks.append(chunk.strip())
            
#             if end >= len(words):
#                 break
            
#             start = end - self.chunk_overlap
        
#         return chunks







# rag/services/document_processor.py
"""
Document Processing Service
Handles file upload, text extraction, chunking, and vectorization.
Supports: PDF, DOCX, TXT, CSV, TSV — including all table formats.

Classes:
    DocumentProcessor    — Main upload handler (Django integration)
    FileTypeDetector     — Auto-detects file type and delimiter
    TabularProcessor     — Converts CSV/TXT rows → natural language
    FreeTextProcessor    — Sliding-window chunker for plain text
    GenericRAGProcessor  — Master router for any file type
    QueryEnhancer        — Fixes vague queries like "this cv"
"""

# ─────────────────────────────────────────────────────────────────────────────
#  IMPORTS
# ─────────────────────────────────────────────────────────────────────────────
import os
import re
import io
import csv as csv_module
import hashlib
import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import Dict, Any, List, Optional

import pandas as pd

from django.conf import settings
from django.core.files.uploadedfile import UploadedFile

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
#  DATA CLASSES
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class Document:
    """A single indexable unit for ChromaDB / any vector store."""
    text:        str
    metadata:    Dict[str, Any]
    doc_id:      str
    source_file: str
    chunk_index: int           = 0
    row_index:   Optional[int] = None


@dataclass
class ProcessingResult:
    """Result returned after processing any file."""
    documents:        List[Document]
    file_type:        str
    total_rows:       int
    columns_detected: List[str]
    warnings:         List[str] = field(default_factory=list)


# ─────────────────────────────────────────────────────────────────────────────
#  1. DOCUMENT PROCESSOR  — Django upload handler
# ─────────────────────────────────────────────────────────────────────────────

class DocumentProcessor:
    """
    Processes uploaded documents for the RAG system.

    KEY DESIGN — Tabular files (CSV, TSV, tabular TXT):
        Each ROW is stored as its own separate chunk with its own embedding.
        This ensures "find Daniel Clark" retrieves EXACTLY his row,
        not a 30-row blob where he is buried.

    Non-tabular files (PDF, DOCX, plain TXT):
        Standard sliding-window chunking.
    """

    def __init__(
        self,
        vector_store,
        embedding_service,
        chunk_size:    int = None,
        chunk_overlap: int = None,
    ):
        self.vector_store      = vector_store
        self.embedding_service = embedding_service
        self.chunk_size        = chunk_size    or getattr(settings, "CHUNK_SIZE",    800)
        self.chunk_overlap     = chunk_overlap or getattr(settings, "CHUNK_OVERLAP", 100)

    # ─────────────────────────────────────────────────────────────────────────
    #  MAIN ENTRY POINT
    # ─────────────────────────────────────────────────────────────────────────

    async def process_document(
        self,
        file:        UploadedFile,
        document_id: str,
    ) -> Dict[str, Any]:
        """
        Full pipeline: detect type → extract → chunk → embed → store.

        Tabular files  → row-by-row indexing (1 chunk per row)
        Other files    → sliding-window chunking
        """
        name_lower = file.name.lower()
        ct         = (file.content_type or "").lower()

        # ── Detect tabular files ──────────────────────────────────────────
        if (
            name_lower.endswith(".csv") or "csv" in ct or
            name_lower.endswith(".tsv") or
            (name_lower.endswith(".txt") and self._peek_is_tabular(file))
        ):
            logger.info(f"[Processor] TABULAR file: '{file.name}' — indexing row by row")
            return await self._process_tabular(file, document_id)
        else:
            logger.info(f"[Processor] TEXT file: '{file.name}' — sliding-window chunks")
            return await self._process_freetext(file, document_id)

    # ─────────────────────────────────────────────────────────────────────────
    #  TABULAR PIPELINE — one chunk per row
    # ─────────────────────────────────────────────────────────────────────────

    async def _process_tabular(
        self,
        file:        UploadedFile,
        document_id: str,
    ) -> Dict[str, Any]:
        """
        Index CSV / TSV / tabular-TXT files row by row.

        Flow:
            Read file → detect delimiter → parse headers + rows
            → convert each row to "Header: Value | Header: Value ..."
            → embed each row separately → store in ChromaDB

        Result:
            100-row CSV → 100 separate chunks, each with its own embedding
            Query "Daniel Clark" → hits exactly his row → correct answer ✅
        """
        content    = file.read()
        raw        = self._decode(content)
        name_lower = file.name.lower()

        # Detect delimiter
        if name_lower.endswith(".tsv"):
            delimiter = "\t"
        elif name_lower.endswith(".csv"):
            delimiter = ","
        else:
            # TXT — detect automatically from content
            lines     = [l for l in raw.splitlines() if l.strip()]
            delimiter = self._detect_delimiter(lines)

        # Parse rows → list of natural-language strings (one per row)
        row_chunks = self._rows_to_nl_list(raw, delimiter)

        if not row_chunks:
            raise ValueError(f"No data rows found in '{file.name}'")

        logger.info(
            f"[Processor] '{file.name}': {len(row_chunks)} rows "
            f"parsed with delimiter '{delimiter}'"
        )

        # Embed all rows in one batch call
        embeddings = self.embedding_service.embed_texts(row_chunks)

        # Build metadata — IMPORTANT: store both 'source' AND 'document_id'
        # 'source'      → used by document_filter (filename-based filtering)
        # 'document_id' → used by document_id-based filtering
        metadatas = [
            {
                "source":       file.name,
                "document_id":  document_id,
                "content_type": file.content_type or "text/csv",
                "chunk_index":  i,
                "chunk_size":   len(chunk),
                "chunk_type":   "table_row",
                "row_number":   i + 1,
            }
            for i, chunk in enumerate(row_chunks)
        ]

        ids = [f"{document_id}_row_{i}" for i in range(len(row_chunks))]

        self.vector_store.add_documents(
            documents  = row_chunks,
            embeddings = embeddings,
            metadata   = metadatas,
            ids        = ids,
        )

        logger.info(f"[Processor] Stored {len(row_chunks)} row-chunks for '{file.name}'")

        return {
            "chunks_created": len(row_chunks),
            "text_length":    sum(len(r) for r in row_chunks),
            "document_id":    document_id,
            "indexing_mode":  "row_by_row",
        }

    # ─────────────────────────────────────────────────────────────────────────
    #  FREE TEXT PIPELINE — sliding window
    # ─────────────────────────────────────────────────────────────────────────

    async def _process_freetext(
        self,
        file:        UploadedFile,
        document_id: str,
    ) -> Dict[str, Any]:
        """Standard pipeline for PDF, DOCX, plain TXT."""
        text = self._extract_text(file)

        if not text or len(text.strip()) < 10:
            raise ValueError(f"No extractable text found in '{file.name}'")

        logger.info(f"[Processor] Extracted {len(text):,} chars from '{file.name}'")

        chunks = self._chunk_text(text)
        if not chunks:
            raise ValueError(f"No chunks generated from '{file.name}'")

        logger.info(f"[Processor] Created {len(chunks)} sliding-window chunks")

        embeddings = self.embedding_service.embed_texts(chunks)

        metadatas = [
            {
                "source":       file.name,
                "document_id":  document_id,
                "content_type": file.content_type or "application/octet-stream",
                "chunk_index":  i,
                "chunk_size":   len(chunk),
                "chunk_type":   "text_chunk",
            }
            for i, chunk in enumerate(chunks)
        ]

        ids = [f"{document_id}_chunk_{i}" for i in range(len(chunks))]

        self.vector_store.add_documents(
            documents  = chunks,
            embeddings = embeddings,
            metadata   = metadatas,
            ids        = ids,
        )

        logger.info(f"[Processor] Stored {len(chunks)} text-chunks for '{file.name}'")

        return {
            "chunks_created": len(chunks),
            "text_length":    len(text),
            "document_id":    document_id,
            "indexing_mode":  "sliding_window",
        }

    # ─────────────────────────────────────────────────────────────────────────
    #  TABULAR HELPERS
    # ─────────────────────────────────────────────────────────────────────────

    def _peek_is_tabular(self, file: UploadedFile) -> bool:
        """
        Peek at first 2 KB of a TXT file to check if it looks tabular.
        Resets file position after reading so the rest of the pipeline
        can read the file normally.
        """
        try:
            pos    = file.tell()
            sample = file.read(2048)
            file.seek(pos)                      # reset — very important!
            raw    = self._decode(sample)
            lines  = [l.strip() for l in raw.splitlines() if l.strip()][:5]
            if not lines:
                return False
            for delim in (",", "\t", ";", "|"):
                counts = [l.count(delim) for l in lines]
                if len(set(counts)) == 1 and counts[0] >= 1:
                    return True
        except Exception:
            pass
        return False

    @staticmethod
    def _detect_delimiter(lines: List[str]) -> str:
        """
        Return the most consistently used delimiter from sample lines.
        Checks: comma, tab, semicolon, pipe.
        """
        best_delim = ","
        best_count = 0
        for delim in (",", "\t", ";", "|"):
            counts = [l.count(delim) for l in lines[:5]]
            if len(set(counts)) == 1 and counts[0] > best_count:
                best_count = counts[0]
                best_delim = delim
        return best_delim

    def _rows_to_nl_list(self, raw: str, delimiter: str = ",") -> List[str]:
        """
        Parse delimited text and return ONE natural-language string PER ROW.

        This is the CRITICAL method — each row gets its own embedding.

        Input (CSV with 3 data rows):
            Student_ID,First_Name,Last_Name,Age,GPA
            S001,Emma,Johnson,16,3.8
            S024,Daniel,Clark,16,3.5
            S003,Olivia,Brown,17,3.9

        Output (list of 3 strings):
            [
              "Student_ID: S001 | First_Name: Emma | Last_Name: Johnson | Age: 16 | GPA: 3.8",
              "Student_ID: S024 | First_Name: Daniel | Last_Name: Clark | Age: 16 | GPA: 3.5",
              "Student_ID: S003 | First_Name: Olivia | Last_Name: Brown | Age: 17 | GPA: 3.9",
            ]

        When user queries "Daniel Clark":
            → embedding matches row 2 precisely ✅
            → NOT buried in 30-row merged blob ✅
        """
        reader = csv_module.reader(io.StringIO(raw), delimiter=delimiter)
        rows   = [r for r in reader if any(c.strip() for c in r)]

        if not rows:
            return []

        # Single row — no header possible
        if len(rows) == 1:
            return [" | ".join(c.strip() for c in rows[0] if c.strip())]

        headers    = [h.strip() for h in rows[0]]
        row_chunks = []

        for row in rows[1:]:
            pairs = []
            for header, cell in zip(headers, row):
                cell = cell.strip()
                if cell:
                    pairs.append(f"{header}: {cell}")

            if pairs:
                row_chunks.append(" | ".join(pairs))

        logger.info(
            f"[Processor] Parsed {len(row_chunks)} rows "
            f"from {len(headers)} columns (delimiter='{delimiter}')"
        )
        return row_chunks

    # ─────────────────────────────────────────────────────────────────────────
    #  TEXT EXTRACTION ROUTER
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_text(self, file: UploadedFile) -> str:
        """Route to the correct extractor based on file type."""
        content    = file.read()
        name_lower = file.name.lower()
        ct         = (file.content_type or "").lower()

        if name_lower.endswith(".pdf") or "pdf" in ct:
            return self._extract_pdf(content)

        if name_lower.endswith(".docx") or "wordprocessingml" in ct:
            return self._extract_docx(content)

        if name_lower.endswith(".txt") or "text/plain" in ct:
            return self._extract_txt(content)

        raise ValueError(
            f"Unsupported file type: '{file.content_type or file.name}'. "
            "Allowed: PDF, DOCX, TXT, CSV, TSV"
        )

    # ─────────────────────────────────────────────────────────────────────────
    #  PDF
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_pdf(self, content: bytes) -> str:
        """
        Extract text and tables from PDF.
        Prefers pdfplumber (table support); falls back to PyPDF2.
        """
        try:
            import pdfplumber
            parts = []

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        for raw_table in (page.extract_tables() or []):
                            if raw_table and len(raw_table) >= 2:
                                table_text = self._table_rows_to_text(raw_table)
                                if table_text:
                                    parts.append(table_text)
                    except Exception as exc:
                        logger.warning(f"[Processor] PDF table page {page_num}: {exc}")

                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            parts.append(page_text.strip())
                    except Exception as exc:
                        logger.warning(f"[Processor] PDF text page {page_num}: {exc}")

            result = "\n\n".join(parts).strip()
            if result:
                return result

        except ImportError:
            logger.warning("[Processor] pdfplumber not found — falling back to PyPDF2")

        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            text   = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text.strip()
        except Exception as exc:
            raise ValueError(f"PDF extraction failed: {exc}") from exc

    # ─────────────────────────────────────────────────────────────────────────
    #  DOCX
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_docx(self, content: bytes) -> str:
        """Extract paragraphs and tables from DOCX in document order."""
        try:
            import docx as docx_lib
            doc   = docx_lib.Document(io.BytesIO(content))
            parts = []
            WNS   = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

            for element in doc.element.body:
                tag = element.tag.split("}")[-1]

                if tag == "p":
                    text = "".join(
                        n.text or "" for n in element.iter()
                        if n.tag.endswith("}t")
                    )
                    if text.strip():
                        parts.append(text.strip())

                elif tag == "tbl":
                    rows = []
                    for tr in element.findall(f".//{WNS}tr"):
                        cells = [
                            "".join(
                                n.text or "" for n in tc.iter()
                                if n.tag.endswith("}t")
                            ).strip()
                            for tc in tr.findall(f".//{WNS}tc")
                        ]
                        if cells:
                            rows.append(cells)
                    if rows:
                        table_text = self._table_rows_to_text(rows)
                        if table_text:
                            parts.append(table_text)

            return "\n\n".join(parts).strip()

        except Exception as exc:
            raise ValueError(f"DOCX extraction failed: {exc}") from exc

    # ─────────────────────────────────────────────────────────────────────────
    #  TXT (plain text only — tabular TXT handled by _process_tabular)
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_txt(self, content: bytes) -> str:
        """
        Extract plain text. Called only for non-tabular TXT files.
        Tabular TXT files are handled by _process_tabular() directly.
        """
        raw = self._decode(content)
        return raw.strip()

    # ─────────────────────────────────────────────────────────────────────────
    #  TABLE → NATURAL LANGUAGE (for PDF/DOCX tables)
    # ─────────────────────────────────────────────────────────────────────────

    def _table_rows_to_text(self, rows: List[List]) -> str:
        """
        Convert list-of-lists table (first row = headers) to
        natural-language strings joined by newlines.

        Used for PDF and DOCX embedded tables.
        [['Name', 'Age'], ['Ahmed', '25']] → "Name: Ahmed | Age: 25"
        """
        if not rows or len(rows) < 2:
            return ""

        cleaned = [
            [str(c).strip() if c is not None else "" for c in row]
            for row in rows
        ]

        headers   = cleaned[0]
        data_rows = [
            row for row in cleaned[1:]
            if any(cell for cell in row)
            and not all(set(cell) <= {"-", "=", "_", " ", ""} for cell in row)
        ]

        if not data_rows:
            return ""

        result = []
        for row in data_rows:
            pairs = [
                f"{(headers[i] if i < len(headers) and headers[i] else f'col_{i}')}: {cell}"
                for i, cell in enumerate(row) if cell
            ]
            if pairs:
                result.append(" | ".join(pairs))

        return "\n".join(result)

    # ─────────────────────────────────────────────────────────────────────────
    #  CHUNKING (for non-tabular files)
    # ─────────────────────────────────────────────────────────────────────────

    def _chunk_text(self, text: str) -> List[str]:
        """Word-based sliding-window chunker for plain text."""
        if not text:
            return []

        words = text.split()

        if len(words) <= self.chunk_size:
            return [" ".join(words)]

        chunks = []
        start  = 0

        while start < len(words):
            end   = min(start + self.chunk_size, len(words))
            chunk = " ".join(words[start:end])
            if chunk.strip():
                chunks.append(chunk.strip())
            if end >= len(words):
                break
            start = end - self.chunk_overlap

        return chunks

    # ─────────────────────────────────────────────────────────────────────────
    #  DECODE HELPER
    # ─────────────────────────────────────────────────────────────────────────

    @staticmethod
    def _decode(file_bytes: bytes) -> str:
        """Try common encodings safely; never crash on bad bytes."""
        for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                return file_bytes.decode(enc)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")


# ─────────────────────────────────────────────────────────────────────────────
#  2. FILE TYPE DETECTOR
# ─────────────────────────────────────────────────────────────────────────────

class FileTypeDetector:
    """Auto-detects file type, delimiter, and encoding. No hardcoding."""

    @staticmethod
    def detect(file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        ext  = path.suffix.lower()

        info = {
            "extension":  ext,
            "file_type":  "unknown",
            "is_tabular": False,
            "delimiter":  None,
            "has_header": False,
            "encoding":   FileTypeDetector._detect_encoding(file_path),
        }

        if ext == ".csv":
            info.update(file_type="csv", is_tabular=True, delimiter=",", has_header=True)
        elif ext == ".tsv":
            info.update(file_type="tsv", is_tabular=True, delimiter="\t", has_header=True)
        elif ext == ".txt":
            info.update(FileTypeDetector._analyze_txt(file_path, info["encoding"]))
        elif ext in {".pdf", ".docx", ".doc"}:
            info.update(file_type=ext.lstrip("."), is_tabular=False)

        return info

    @staticmethod
    def _detect_encoding(file_path: str) -> str:
        for enc in ("utf-8", "latin-1", "utf-16", "cp1252"):
            try:
                with open(file_path, "r", encoding=enc) as f:
                    f.read(1024)
                return enc
            except (UnicodeDecodeError, OSError):
                continue
        return "latin-1"

    @staticmethod
    def _analyze_txt(file_path: str, encoding: str) -> Dict[str, Any]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                sample = [f.readline() for _ in range(10)]
            sample = [l.strip() for l in sample if l.strip()]
        except Exception:
            return {"file_type": "txt_freetext", "is_tabular": False}

        if not sample:
            return {"file_type": "txt_freetext", "is_tabular": False}

        counts = {d: sum(l.count(d) for l in sample) for d in (",", "\t", "|", ";")}
        best   = max(counts, key=counts.get)

        if counts[best] / max(len(sample), 1) >= 1:
            field_counts  = [len(l.split(best)) for l in sample]
            is_consistent = len(set(field_counts)) <= 2
            if is_consistent:
                first_fields = sample[0].split(best)
                has_header   = all(
                    not re.match(r"^[0-9.]+$", f.strip())
                    for f in first_fields[:5]
                )
                return {
                    "file_type":  "txt_tabular",
                    "is_tabular": True,
                    "delimiter":  best,
                    "has_header": has_header,
                }

        return {"file_type": "txt_freetext", "is_tabular": False}


# ─────────────────────────────────────────────────────────────────────────────
#  3. TABULAR PROCESSOR
# ─────────────────────────────────────────────────────────────────────────────

class TabularProcessor:
    """
    Converts ANY tabular file into RAG-ready Documents.
    Column names auto-detected — never hardcoded.
    Each row = one Document with its own embedding.
    """

    def process(self, file_path: str, file_info: Dict) -> ProcessingResult:
        try:
            df = self._load_dataframe(file_path, file_info)
        except Exception as exc:
            return ProcessingResult(
                documents=[], file_type=file_info["file_type"],
                total_rows=0, columns_detected=[],
                warnings=[f"Failed to load: {exc}"],
            )

        if df.empty:
            return ProcessingResult(
                documents=[], file_type=file_info["file_type"],
                total_rows=0, columns_detected=[],
                warnings=["File is empty"],
            )

        df.columns = [self._clean_col(c) for c in df.columns]
        df         = df.dropna(how="all")
        col_types  = self._detect_col_types(df)
        filename   = Path(file_path).name
        documents  = []

        for idx, row in df.iterrows():
            documents.append(self._row_to_document(row, idx, col_types, filename))

        if len(df) > 10:
            documents.extend(self._summary_chunks(df, filename, col_types))

        warnings = []
        if len(df) > 1000:
            warnings.append(f"Large file ({len(df)} rows) — consider batch indexing.")

        logger.info(f"[TabularProcessor] '{filename}' → {len(documents)} documents")

        return ProcessingResult(
            documents        = documents,
            file_type        = file_info["file_type"],
            total_rows       = len(df),
            columns_detected = list(df.columns),
            warnings         = warnings,
        )

    def _load_dataframe(self, file_path: str, file_info: Dict) -> "pd.DataFrame":
        enc   = file_info.get("encoding", "utf-8")
        delim = file_info.get("delimiter", ",")
        ext   = file_info.get("extension", "")

        if ext == ".csv":
            try:
                return pd.read_csv(file_path, encoding=enc, dtype=str)
            except Exception:
                return pd.read_csv(file_path, encoding=enc, sep=None, engine="python", dtype=str)

        if file_info["file_type"] == "txt_tabular":
            return pd.read_csv(
                file_path, sep=delim, encoding=enc, dtype=str,
                header=0 if file_info.get("has_header") else None,
                on_bad_lines="skip",
            )

        return pd.DataFrame()

    @staticmethod
    def _clean_col(col: str) -> str:
        col = re.sub(r"[_\-]+", " ", str(col).strip())
        return re.sub(r"\s+", " ", col).strip().title()

    @staticmethod
    def _detect_col_types(df: "pd.DataFrame") -> Dict[str, str]:
        types = {}
        for col in df.columns:
            sample = df[col].dropna().head(10)
            n      = max(len(sample), 1)

            numeric_ratio = sum(
                1 for v in sample if re.match(r"^-?[\d.]+$", str(v).strip())
            ) / n

            if numeric_ratio >= 0.8:
                col_lower = col.lower()
                if any(kw in col_lower for kw in ("id", "code", "no", "number")):
                    types[col] = "identifier"
                elif sum(1 for v in sample if re.match(r"^(19|20)\d{2}$", str(v).strip())) / n >= 0.8:
                    types[col] = "year"
                else:
                    types[col] = "numeric"
            elif sum(1 for v in sample if "@" in str(v)) / n >= 0.5:
                types[col] = "email"
            else:
                types[col] = "text"

        return types

    def _row_to_document(
        self,
        row:       "pd.Series",
        row_index: int,
        col_types: Dict[str, str],
        filename:  str,
    ) -> Document:
        record_type = self._record_type(filename)
        parts       = []
        metadata    = {
            "source_file": filename,
            "row_index":   row_index,
            "record_type": record_type,
        }

        for col in row.index:
            val = row[col]
            if pd.isna(val) or str(val).strip() in ("", "nan", "None"):
                continue
            val = str(val).strip()
            parts.append(f"{col}: {val}")
            metadata[col.lower().replace(" ", "_")] = val

        fields_text = " | ".join(parts)
        name        = self._extract_name(row)
        text        = (
            f"{name} — {record_type} from {filename}:\n{fields_text}"
            if name else
            f"{record_type} from {filename}:\n{fields_text}"
        )

        doc_id = hashlib.md5(f"{filename}_{row_index}".encode()).hexdigest()[:16]

        return Document(
            text=text, metadata=metadata, doc_id=doc_id,
            source_file=filename, row_index=row_index,
        )

    def _summary_chunks(
        self,
        df:        "pd.DataFrame",
        filename:  str,
        col_types: Dict[str, str],
    ) -> List[Document]:
        docs        = []
        record_type = self._record_type(filename)
        col_list    = ", ".join(df.columns)

        overview = (
            f"Overview of {filename}: Contains {len(df)} {record_type}s. "
            f"Columns: {col_list}. "
        )

        for col, ctype in col_types.items():
            if ctype == "numeric":
                try:
                    vals = pd.to_numeric(df[col], errors="coerce").dropna()
                    if len(vals):
                        overview += (
                            f"{col} — min: {vals.min():.2f}, "
                            f"max: {vals.max():.2f}, avg: {vals.mean():.2f}. "
                        )
                except Exception:
                    pass

        docs.append(Document(
            text=overview,
            metadata={"source_file": filename, "chunk_type": "summary"},
            doc_id=hashlib.md5(f"{filename}_overview".encode()).hexdigest()[:16],
            source_file=filename,
        ))

        for col, ctype in col_types.items():
            if ctype == "text" and 1 < df[col].nunique() < 30:
                unique_vals = df[col].dropna().unique().tolist()
                vals_str    = ", ".join(str(v) for v in unique_vals[:25])
                cat_text    = (
                    f"All unique values for '{col}' in {filename}: {vals_str}. "
                    f"Total unique: {len(unique_vals)}."
                )
                docs.append(Document(
                    text=cat_text,
                    metadata={"source_file": filename, "chunk_type": "categorical", "column": col},
                    doc_id=hashlib.md5(f"{filename}_{col}_cat".encode()).hexdigest()[:16],
                    source_file=filename,
                ))

        return docs

    @staticmethod
    def _record_type(filename: str) -> str:
        stem = Path(filename).stem.lower()
        stem = re.sub(r"[_\-\.]", " ", stem)
        stem = re.sub(r"\d+", "", stem)
        stem = re.sub(r"\b(data|file|report|sheet|export|list)\b", "", stem)
        stem = re.sub(r"\s+", " ", stem).strip()
        if not stem:
            return "Record"
        if stem.endswith("s") and len(stem) > 4:
            stem = stem[:-1]
        return stem.title() + " record"

    @staticmethod
    def _extract_name(row: "pd.Series") -> Optional[str]:
        row_lower = {k.lower(): v for k, v in row.items()}

        for key in ("name", "full name", "fullname", "title"):
            val = str(row_lower.get(key, "")).strip()
            if val and val not in ("nan", "None", ""):
                return val

        first = str(row_lower.get("first name", row_lower.get("firstname", ""))).strip()
        last  = str(row_lower.get("last name",  row_lower.get("lastname",  ""))).strip()
        if first and last and first != "nan" and last != "nan":
            return f"{first} {last}"

        return None


# ─────────────────────────────────────────────────────────────────────────────
#  4. FREE TEXT PROCESSOR
# ─────────────────────────────────────────────────────────────────────────────

class FreeTextProcessor:
    """Sliding-window chunker for non-tabular text files."""

    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap    = overlap

    def process(self, file_path: str, file_info: Dict) -> ProcessingResult:
        enc      = file_info.get("encoding", "utf-8")
        filename = Path(file_path).name

        try:
            with open(file_path, "r", encoding=enc) as f:
                text = f.read()
        except Exception as exc:
            return ProcessingResult(
                documents=[], file_type="txt_freetext",
                total_rows=0, columns_detected=[],
                warnings=[f"Could not read: {exc}"],
            )

        chunks    = self._split(text)
        documents = []

        for i, chunk in enumerate(chunks):
            doc_id = hashlib.md5(f"{filename}_{i}".encode()).hexdigest()[:16]
            documents.append(Document(
                text=chunk.strip(),
                metadata={"source_file": filename, "chunk_index": i, "chunk_type": "text"},
                doc_id=doc_id,
                source_file=filename,
                chunk_index=i,
            ))

        logger.info(f"[FreeTextProcessor] '{filename}' → {len(documents)} chunks")

        return ProcessingResult(
            documents=documents, file_type="txt_freetext",
            total_rows=len(chunks), columns_detected=[],
        )

    def _split(self, text: str) -> List[str]:
        chunks = []
        start  = 0

        while start < len(text):
            end   = start + self.chunk_size
            chunk = text[start:end]

            if end < len(text):
                boundary = chunk.rfind(". ")
                if boundary > self.chunk_size * 0.5:
                    chunk = chunk[: boundary + 1]
                    end   = start + boundary + 1

            if chunk.strip():
                chunks.append(chunk)

            start = end - self.overlap

        return chunks


# ─────────────────────────────────────────────────────────────────────────────
#  5. GENERIC RAG PROCESSOR
# ─────────────────────────────────────────────────────────────────────────────

class GenericRAGProcessor:
    """Master router — give it any file, get back RAG-ready Documents."""

    def __init__(self, chunk_size: int = 500):
        self.tabular_processor = TabularProcessor()
        self.text_processor    = FreeTextProcessor(chunk_size=chunk_size)

    def process(self, file_path: str) -> ProcessingResult:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_info = FileTypeDetector.detect(file_path)

        logger.info(
            f"[GenericRAGProcessor] '{Path(file_path).name}' | "
            f"type={file_info['file_type']} | tabular={file_info['is_tabular']}"
        )

        if file_info["is_tabular"]:
            return self.tabular_processor.process(file_path, file_info)
        else:
            return self.text_processor.process(file_path, file_info)

    def process_folder(self, folder_path: str) -> List[ProcessingResult]:
        supported = {".csv", ".tsv", ".txt", ".md"}
        results   = []
        for f in Path(folder_path).iterdir():
            if f.suffix.lower() in supported:
                results.append(self.process(str(f)))
        return results


# ─────────────────────────────────────────────────────────────────────────────
#  6. QUERY ENHANCER
# ─────────────────────────────────────────────────────────────────────────────

class QueryEnhancer:
    """
    Fixes vague queries before they reach the retriever.

    "skills of this cv"  +  active_file="Syed_Shahzad_Ali.pdf"
    → "skills of Syed Shahzad Ali"
    """

    WEAK_PHRASES = [
        "this document", "this file",   "this cv",     "this resume",
        "this data",     "this table",  "this report", "this record",
        "the document",  "the file",    "the data",    "the cv",
        "the resume",    "the report",  "this pdf",    "the pdf",
    ]

    @staticmethod
    def enhance(query: str, active_file: Optional[str] = None) -> str:
        if not active_file:
            return query

        label    = Path(active_file).stem.replace("_", " ").replace("-", " ")
        enhanced = query

        for phrase in QueryEnhancer.WEAK_PHRASES:
            if phrase.lower() in enhanced.lower():
                enhanced = re.sub(
                    re.escape(phrase), label, enhanced, flags=re.IGNORECASE,
                )
                break

        if enhanced != query:
            logger.info(f"[QueryEnhancer] '{query}' → '{enhanced}'")

        return enhanced

    @staticmethod
    def fallback_keyword_search(
        query:       str,
        collection,
        top_k:       int           = 5,
        filter_file: Optional[str] = None,
    ) -> List[Dict]:
        keywords = re.findall(r"\b[A-Z][a-z]+\b", query)

        if not keywords:
            logger.warning("[QueryEnhancer] Fallback: no keywords found")
            return []

        logger.info(f"[QueryEnhancer] Fallback keywords: {keywords}")
        where = {"source": {"$eq": filter_file}} if filter_file else None

        for kw in keywords:
            try:
                results = collection.query(
                    query_texts=[kw], n_results=top_k, where=where,
                    include=["documents", "metadatas", "distances"],
                )
                if results and results["documents"] and results["documents"][0]:
                    logger.info(f"[QueryEnhancer] Fallback hit: '{kw}'")
                    return [
                        {
                            "text":            doc,
                            "metadata":        meta,
                            "relevance_score": round(1 - dist, 4),
                        }
                        for doc, meta, dist in zip(
                            results["documents"][0],
                            results["metadatas"][0],
                            results["distances"][0],
                        )
                    ]
            except Exception as exc:
                logger.warning(f"[QueryEnhancer] Fallback error '{kw}': {exc}")
                continue

        logger.warning("[QueryEnhancer] Fallback: no results")
        return []